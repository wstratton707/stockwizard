"""portfolio_docx.py — the Word (.docx) portfolio review.

A written client report on a tracked portfolio: executive summary, allocation,
holdings, risk, performance attribution, strengths and concerns, investment
insights and considerations, with an appendix carrying the full holdings table
and the methodology behind every number.

Reuses `portfolio_excel._derive_portfolio_stats` for the analytics and
`portfolio_narrative.build_narrative` for the prose, so the Word report, the
PowerPoint deck and the Excel workbook can never disagree about the same
portfolio. Formatting helpers come from `docx_builder` for one visual language
across every document the app produces.

    from portfolio_docx import build_portfolio_docx
    buf = build_portfolio_docx(name, tracked, profiles)
"""

from __future__ import annotations

import io
import os
from datetime import date

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from docx_builder import (NAVY, BLUE, SLATE, GREEN, RED, MUTED, DISCLAIMER,
                          _ASSET_LOGO, _heading, _para, _bullet, _table,
                          _kv_table)
from portfolio_excel import _derive_portfolio_stats
from portfolio_narrative import build_narrative

# Matching the on-screen corridor blue so charts read as the same product.
C_BLUE, C_LIGHT = "#3F6C9C", "#C3DDF5"
C_GREEN, C_RED, C_GREY = "#15803d", "#b91c1c", "#94a3b8"
SECTOR_PALETTE = ["#3F6C9C", "#6C8FB8", "#9BB4D2", "#C3DDF5", "#D98324",
                  "#E0B341", "#7F9A6B", "#8E7CC3", "#B0BEC5", "#5D8AA8"]


# ── low-level Word plumbing python-docx has no API for ────────────────────────
def _field(paragraph, instr):
    """Insert a Word field code (used for TOC and PAGE numbers)."""
    r = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar"); fld_begin.set(qn("w:fldCharType"), "begin")
    instr_el = OxmlElement("w:instrText"); instr_el.set(qn("xml:space"), "preserve")
    instr_el.text = instr
    fld_sep = OxmlElement("w:fldChar"); fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar"); fld_end.set(qn("w:fldCharType"), "end")
    r._r.append(fld_begin); r._r.append(instr_el)
    r._r.append(fld_sep); r._r.append(fld_end)
    return r


def _toc(doc):
    """Table-of-contents field. Word renders it once the reader updates fields;
    until then the placeholder line below explains why it looks empty."""
    p = doc.add_paragraph()
    _field(p, r'TOC \o "1-2" \h \z \u')


def _page_footer(section, left_text):
    """Footer with the brand line on the left and 'Page X of Y' on the right."""
    f = section.footer
    p = f.paragraphs[0] if f.paragraphs else f.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in list(p.runs):
        run.text = ""
    r = p.add_run(left_text + "\t\t")
    r.font.size = Pt(8)
    r.font.color.rgb = MUTED
    r2 = p.add_run("Page ")
    r2.font.size = Pt(8); r2.font.color.rgb = MUTED
    _field(p, "PAGE")
    r3 = p.add_run(" of ")
    r3.font.size = Pt(8); r3.font.color.rgb = MUTED
    _field(p, "NUMPAGES")


def _page_header(section, text):
    h = section.header
    p = h.paragraphs[0] if h.paragraphs else h.add_paragraph()
    for run in list(p.runs):
        run.text = ""
    r = p.add_run(text)
    r.font.size = Pt(8)
    r.font.color.rgb = MUTED
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def _h1(doc, text):
    """Heading that the TOC field can actually see (uses Word's built-in style)."""
    p = doc.add_heading(text, level=1)
    for r in p.runs:
        r.font.color.rgb = NAVY
        r.font.size = Pt(15)
        r.font.name = "Calibri"
    return p


def _h2(doc, text):
    p = doc.add_heading(text, level=2)
    for r in p.runs:
        r.font.color.rgb = BLUE
        r.font.size = Pt(11.5)
        r.font.name = "Calibri"
    return p


def _money(v, dp=0):
    """Currency with the sign OUTSIDE the symbol: -$67, never $-67."""
    if v is None:
        return "—"
    return f"-${abs(v):,.{dp}f}" if v < 0 else f"${v:,.{dp}f}"


def _clip(s, n):
    """Truncate on a word boundary so names don't break mid-word."""
    s = (s or "").strip()
    if len(s) <= n:
        return s
    cut = s[:n].rsplit(" ", 1)[0]
    return (cut if len(cut) >= n * 0.6 else s[:n]).rstrip(",.") + "…"


def _pc(v, dp=1, sign=False):
    if v is None:
        return "—"
    return f"{v:+.{dp}f}%" if sign else f"{v:.{dp}f}%"


def _w(v):
    return "—" if v is None else f"{v:.1%}"


# ── charts ────────────────────────────────────────────────────────────────────
def _fig_buf(fig):
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=170, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    buf.seek(0)
    return buf


def _alloc_pie(stats, by="sectors", title="Allocation"):
    items = stats.get(by) or []
    total = stats.get("total_value") or 1
    if not items:
        return None
    top = items[:8]
    rest = items[8:]
    labels = [k for k, _ in top]
    vals = [v for _, v in top]
    if rest:
        labels.append(f"Other ({len(rest)})")
        vals.append(sum(v for _, v in rest))
    try:
        fig, ax = plt.subplots(figsize=(5.4, 3.4))
        wedges, *_ = ax.pie(
            vals, startangle=90, counterclock=False,
            colors=SECTOR_PALETTE[:len(vals)],
            wedgeprops=dict(width=0.42, edgecolor="white", linewidth=1.4))
        ax.legend(wedges, [f"{l} — {v/total:.1%}" for l, v in zip(labels, vals)],
                  loc="center left", bbox_to_anchor=(1.0, 0.5),
                  fontsize=8, frameon=False)
        ax.set_title(title, fontsize=10.5, color="#0f172a", loc="left", pad=10)
        ax.axis("equal")
        return _fig_buf(fig)
    except Exception:
        plt.close("all")
        return None


def _weights_bar(stats):
    rows = stats.get("rows") or []
    if not rows:
        return None
    top = rows[:12][::-1]
    try:
        fig, ax = plt.subplots(figsize=(6.6, max(2.4, 0.30 * len(top))))
        ax.barh([r["ticker"] for r in top], [r["w"] * 100 for r in top],
                color=C_BLUE, height=0.68)
        for i, r in enumerate(top):
            ax.text(r["w"] * 100 + 0.35, i, f"{r['w']:.1%}",
                    va="center", fontsize=7.5, color="#475569")
        ax.set_xlabel("Portfolio weight (%)", fontsize=8, color="#475569")
        ax.grid(axis="x", alpha=0.25)
        ax.set_axisbelow(True)
        for s in ("top", "right"):
            ax.spines[s].set_visible(False)
        ax.tick_params(labelsize=8, colors="#475569")
        ax.set_title("Position weights", fontsize=10.5, color="#0f172a", loc="left")
        return _fig_buf(fig)
    except Exception:
        plt.close("all")
        return None


def _value_chart(tracked):
    curve = tracked.get("curve")
    if curve is None or len(curve) < 2:
        return None
    try:
        fig, ax = plt.subplots(figsize=(6.8, 2.9))
        ax.plot(curve.index, curve["Portfolio"], color=C_BLUE, lw=1.7,
                label="Portfolio", zorder=3)
        if "SP500" in curve.columns and not curve["SP500"].isna().all():
            ax.plot(curve.index, curve["SP500"], color=C_GREY, lw=1.2, ls="--",
                    label="S&P 500 (same contributions)")
        if "Contrib" in curve.columns:
            ax.plot(curve.index, curve["Contrib"], color="#D98324", lw=1.1, ls=":",
                    label="Capital invested")
        ax.set_title("Portfolio value since inception", fontsize=10.5,
                     color="#0f172a", loc="left")
        ax.legend(fontsize=7.5, frameon=False, loc="upper left")
        ax.grid(alpha=0.22)
        for s in ("top", "right"):
            ax.spines[s].set_visible(False)
        ax.tick_params(labelsize=7.5, colors="#64748b")
        ax.yaxis.set_major_formatter(plt.FuncFormatter(lambda v, _: f"${v:,.0f}"))
        return _fig_buf(fig)
    except Exception:
        plt.close("all")
        return None


def _contrib_bar(narr):
    cs = (narr.get("winners") or []) + (narr.get("losers") or [])
    if not cs:
        return None
    cs = sorted(cs, key=lambda c: c["gain_usd"])
    try:
        fig, ax = plt.subplots(figsize=(6.6, max(2.2, 0.30 * len(cs))))
        ax.barh([c["ticker"] for c in cs], [c["gain_usd"] for c in cs],
                color=[C_GREEN if c["gain_usd"] >= 0 else C_RED for c in cs],
                height=0.66)
        ax.axvline(0, color="#94a3b8", lw=0.8)
        ax.set_xlabel("Gain / loss vs cost basis ($)", fontsize=8, color="#475569")
        ax.grid(axis="x", alpha=0.25)
        ax.set_axisbelow(True)
        for s in ("top", "right"):
            ax.spines[s].set_visible(False)
        ax.tick_params(labelsize=8, colors="#475569")
        ax.set_title("Contribution by holding", fontsize=10.5, color="#0f172a",
                     loc="left")
        return _fig_buf(fig)
    except Exception:
        plt.close("all")
        return None


def _pic(doc, buf, width=6.4):
    if buf is None:
        return
    try:
        doc.add_picture(buf, width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception:
        pass


# ── the report ────────────────────────────────────────────────────────────────
def build_portfolio_docx(portfolio_name, tracked, profiles):
    """Return a BytesIO .docx portfolio review."""
    stats = _derive_portfolio_stats(tracked, profiles)
    narr = build_narrative(portfolio_name, tracked, stats)
    m = tracked.get("metrics", {}) or {}
    today = date.today().strftime("%B %d, %Y")
    inception = tracked.get("inception_date", "—")

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

    sec = doc.sections[0]
    sec.left_margin = sec.right_margin = Inches(0.85)
    sec.top_margin = sec.bottom_margin = Inches(0.75)
    _page_header(sec, f"{portfolio_name} · Portfolio Review · {today}")
    _page_footer(sec, "QuantWizard · For informational purposes only. Not investment advice.")

    # ── Cover ─────────────────────────────────────────────────────────────────
    if os.path.exists(_ASSET_LOGO):
        try:
            doc.add_picture(_ASSET_LOGO, width=Inches(1.9))
        except Exception:
            pass
    t = doc.add_paragraph()
    t.paragraph_format.space_before = Pt(60)
    r = t.add_run("Portfolio Review")
    r.bold = True; r.font.size = Pt(30); r.font.color.rgb = NAVY
    s = doc.add_paragraph()
    r = s.add_run(portfolio_name)
    r.font.size = Pt(17); r.font.color.rgb = BLUE
    _para(doc, f"Prepared {today}   ·   Tracked since {inception}", size=10,
          color=MUTED, italic=True, after=2)

    _cur = m.get("Final Value")
    _gain = m.get("Total Gain/Loss")
    _tr = m.get("Total Return")
    hero = doc.add_paragraph()
    hero.paragraph_format.space_before = Pt(26)
    r = hero.add_run(_money(_cur))
    r.bold = True; r.font.size = Pt(26)
    r.font.color.rgb = NAVY
    r2 = hero.add_run(f"   {_pc(_tr, sign=True)} since inception")
    r2.font.size = Pt(12)
    r2.font.color.rgb = GREEN if (_tr or 0) >= 0 else RED
    _para(doc, narr["objective"], size=11, color=SLATE, after=14)

    _para(doc, "Contents", size=12, bold=True, color=NAVY, after=2)
    _para(doc, "If the list below is blank, press Ctrl+A then F9 in Word to build it.",
          size=8, color=MUTED, italic=True, after=4)
    _toc(doc)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # ── 1. Executive summary ──────────────────────────────────────────────────
    _h1(doc, "Executive Summary")
    _kv_table(doc, [
        ("Portfolio", portfolio_name),
        ("Report date", today),
        ("Tracking since", inception),
        ("Current market value", _money(_cur)),
        ("Capital invested", _money(m.get("Total Contributed"))),
        ("Total gain / loss", f"{_money(_gain)}  ({_pc(_tr, sign=True)})"),
        ("Number of holdings", stats["n"]),
        ("Largest position", (f"{stats['largest']['ticker']} · {_w(stats['largest']['w'])}"
                              if stats.get("largest") else "—")),
        ("Sectors represented", len([1 for k, _ in stats["sectors"] if k != "Unknown"])),
        ("Portfolio beta", "—" if stats["beta"] is None else f"{stats['beta']:.2f}"),
        ("Annualised volatility", _pc(stats.get("vol"))),
        ("Dividend yield", _w(stats.get("div_yield"))),
    ])
    _h2(doc, "Objective")
    _para(doc, narr["objective"])
    _h2(doc, "Overall assessment")
    _para(doc, narr["assessment"])

    # ── 2. Asset allocation ───────────────────────────────────────────────────
    _h1(doc, "Asset Allocation Analysis")
    _para(doc, narr["sector_commentary"])
    _pic(doc, _alloc_pie(stats, "sectors", "Allocation by sector"), 5.9)
    _pic(doc, _weights_bar(stats), 6.2)

    _h2(doc, "Allocation by sector")
    _tot = stats["total_value"] or 1
    _table(doc, ["Sector", "Value", "Weight"],
           [[s, _money(v), _w(v / _tot)] for s, v in stats["sectors"]])
    _h2(doc, "Allocation by industry")
    _table(doc, ["Industry", "Value", "Weight"],
           [[s, _money(v), _w(v / _tot)] for s, v in stats["industries"][:12]])
    _h2(doc, "Concentration")
    _kv_table(doc, [
        ("Largest position", (f"{stats['largest']['ticker']} · {_w(stats['largest']['w'])}"
                              if stats.get("largest") else "—")),
        ("Top 5 holdings", _w(stats["top5"])),
        ("Top 10 holdings", _w(stats["top10"])),
        ("Herfindahl index (HHI)", f"{stats['hhi']:.3f}"),
        ("Effective number of holdings",
         "—" if not stats.get("eff_n") else f"{stats['eff_n']:.1f}"),
    ])

    # ── 3. Holdings detail ────────────────────────────────────────────────────
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    _h1(doc, "Portfolio Holdings Detail")
    _para(doc, "Sorted by portfolio weight. Cost basis is what was paid for the "
               "shares still held; unrealised gain is measured against it.",
          size=9.5, color=MUTED)
    _table(doc,
           ["Ticker", "Company", "Shares", "Price", "Value", "Weight",
            "Cost basis", "Unrealised", "Return"],
           [[r["ticker"], _clip(r["name"], 26), f"{r['shares']:,.3f}",
             _money(r["last_price"], 2), _money(r["value"]), _w(r["w"]),
             _money(r.get("cost_basis")),
             _money((r["value"] - r["cost_basis"]) if r.get("cost_basis") else None),
             _pc(r.get("gain_pct"), sign=True)]
            for r in stats["rows"]])

    # ── 4. Risk ───────────────────────────────────────────────────────────────
    _h1(doc, "Risk Analysis")
    # With only a few sessions, a drawdown of 0.0% and a Sharpe of 0 are what
    # the formulas return, not what the portfolio did. Printing them as figures
    # gives false comfort, so they are withheld until there is enough history.
    _thin = stats.get("thin_history")
    _hold = "Needs ~1 month of history"
    _kv_table(doc, [
        ("Portfolio beta", "—" if stats["beta"] is None else f"{stats['beta']:.2f}"),
        ("Annualised volatility", _hold if _thin else _pc(stats.get("vol"))),
        ("Maximum drawdown", _hold if _thin else _pc(m.get("Max Drawdown"))),
        ("Sharpe ratio", _hold if _thin else m.get("Sharpe Ratio", "—")),
        ("Sortino ratio", _hold if _thin else m.get("Sortino Ratio", "—")),
        ("Concentration (HHI)", f"{stats['hhi']:.3f} — {narr['diversification_band']}"),
        ("Diversification score", f"{narr['diversification_score']} / 10"),
    ])
    if _thin:
        _para(doc, f"Risk statistics are withheld above because only "
                   f"{stats['n_days']} trading days have elapsed since inception. "
                   f"A volatility or drawdown figure computed from that window "
                   f"describes a couple of sessions, not the portfolio.",
              size=9.5, color=MUTED, italic=True)
    _h2(doc, "What the risk profile means")
    _para(doc, narr["posture_text"])
    _para(doc, narr["diversification_text"])
    _h2(doc, "Risk observations")
    for c in narr["concerns"]:
        _bullet(doc, c)

    # ── 5. Performance ────────────────────────────────────────────────────────
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    _h1(doc, "Performance Analysis")
    _pic(doc, _value_chart(tracked), 6.4)
    _kv_table(doc, [
        ("Total return", _pc(m.get("Total Return"), sign=True)),
        ("S&P 500, same contributions", _pc(m.get("S&P 500 Return"), sign=True)
         if isinstance(m.get("S&P 500 Return"), (int, float)) else "—"),
        ("Relative performance",
         (f"{m['vs S&P 500']:+.1f} pts" if isinstance(m.get("vs S&P 500"), (int, float))
          else "—")),
        ("Annualised return", _pc(m.get("Ann. Return"), sign=True)
         if m.get("Ann. Return") is not None else "Needs 3 months of history"),
        # A portfolio two days old has one "month", so best = worst and
        # "100% positive months" is arithmetic, not a track record.
        ("Best month", _hold if _thin else _pc(m.get("Best Month"), sign=True)),
        ("Worst month", _hold if _thin else _pc(m.get("Worst Month"), sign=True)),
        ("Positive months", _hold if _thin else _pc(m.get("% Months Positive"))),
    ])
    for line in narr["performance_lines"]:
        _para(doc, line)
    _pic(doc, _contrib_bar(narr), 6.2)
    if narr["winners"]:
        _h2(doc, "Largest contributors")
        _table(doc, ["Ticker", "Company", "Gain", "Return", "Weight"],
               [[c["ticker"], _clip(c["name"], 26), _money(c["gain_usd"]),
                 _pc(c.get("gain_pct"), sign=True), _w(c.get("weight"))]
                for c in narr["winners"]])
    if narr["losers"]:
        _h2(doc, "Largest detractors")
        _table(doc, ["Ticker", "Company", "Loss", "Return", "Weight"],
               [[c["ticker"], _clip(c["name"], 26), _money(c["gain_usd"]),
                 _pc(c.get("gain_pct"), sign=True), _w(c.get("weight"))]
                for c in narr["losers"]])

    # ── 6. Strengths & weaknesses ─────────────────────────────────────────────
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    _h1(doc, "Portfolio Strengths & Weaknesses")
    _h2(doc, "Key strengths")
    for s_ in narr["strengths"]:
        _bullet(doc, s_)
    _h2(doc, "Potential concerns")
    for c in narr["concerns"]:
        _bullet(doc, c)
    _h2(doc, "Diversification assessment")
    _para(doc, narr["diversification_text"])
    _h2(doc, "Sector exposure")
    _para(doc, narr["sector_commentary"])

    # ── 7. Investment insights ────────────────────────────────────────────────
    _h1(doc, "Investment Insights")
    _h2(doc, "What drives this portfolio")
    for t_ in narr["themes"]:
        _bullet(doc, t_)
    _h2(doc, f"Style: {narr['style_label']}")
    _para(doc, narr["style_text"])
    _h2(doc, f"Positioning: {narr['posture_label']}")
    _para(doc, narr["posture_text"])
    _h2(doc, "Where this portfolio should do well")
    for e in narr["environments_favourable"]:
        _bullet(doc, e[0].upper() + e[1:])
    _h2(doc, "Where it is likely to struggle")
    for e in narr["environments_adverse"]:
        _bullet(doc, e[0].upper() + e[1:])

    # ── 8. Recommendations ────────────────────────────────────────────────────
    _h1(doc, "Considerations")
    _para(doc, "Observations that follow from the portfolio's current structure. "
               "These are not recommendations to buy or sell — they are the points "
               "a reviewer would raise.", size=9.5, color=MUTED)
    for rec in narr["recommendations"]:
        _bullet(doc, rec)

    # ── 9. Appendix ───────────────────────────────────────────────────────────
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    _h1(doc, "Appendix")
    _h2(doc, "Complete holdings")
    _table(doc,
           ["Ticker", "Company", "Sector", "Industry", "Shares", "Value",
            "Weight", "P/E", "Beta", "Yield"],
           [[r["ticker"], _clip(r["name"], 22), _clip(r["sector"], 18),
             _clip(r["industry"], 18), f"{r['shares']:,.2f}", _money(r["value"]),
             _w(r["w"]),
             "—" if r.get("pe") is None else f"{r['pe']:.1f}",
             "—" if r.get("beta") is None else f"{r['beta']:.2f}",
             _w(r.get("div_yield"))]
            for r in stats["rows"]])

    _h2(doc, "Portfolio statistics")
    _kv_table(doc, [
        ("Holdings", stats["n"]),
        ("Market value", _money(stats["total_value"])),
        ("Weighted P/E", "—" if stats.get("wpe") is None else f"{stats['wpe']:.1f}x"),
        ("Weighted revenue growth", _w(stats.get("rev_growth"))),
        ("Weighted earnings growth", _w(stats.get("eps_growth"))),
        ("Dividend yield", _w(stats.get("div_yield"))),
        ("Portfolio beta", "—" if stats["beta"] is None else f"{stats['beta']:.2f}"),
        ("Trading days tracked", stats["n_days"]),
    ])

    _h2(doc, "Methodology")
    for line in [
        "Performance is measured forward from the day each holding was added — "
        "this is a mark-to-market record, not a backtest of a strategy.",
        "Returns are time-weighted, so adding or withdrawing money never counts "
        "as investment performance.",
        "The S&P 500 comparison funds an SPY position on the same dates and in "
        "the same amounts, making it a like-for-like benchmark rather than a "
        "lump-sum one.",
        "Portfolio beta is the weighted average of holdings' betas; positions "
        "without a published beta are treated as 1.0.",
        "Weighted P/E is earnings-weighted (harmonic) across holdings with "
        "positive earnings, so one richly-valued position cannot dominate it.",
        "Concentration uses the Herfindahl index (sum of squared weights). The "
        "effective number of holdings is its reciprocal — how many equally-sized "
        "positions would carry the same concentration.",
        "Dividend yield counts non-payers as zero, because they are part of the "
        "portfolio the yield is quoted on.",
        "Contribution is measured in dollars against cost basis rather than as a "
        "percentage return, so position size is reflected in the ranking.",
    ]:
        _bullet(doc, line)

    _h2(doc, "Data sources")
    _para(doc, "Prices and company fundamentals: Yahoo Finance. Corporate filings: "
               "SEC EDGAR. Company reference data: Polygon.io. Figures are as of "
               f"{today} and may be delayed or unavailable for some securities.",
          size=9.5)

    _para(doc, DISCLAIMER, size=8, color=MUTED, italic=True, after=0)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
